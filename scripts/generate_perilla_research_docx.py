#!/usr/bin/env python3
"""生成 Perilla 技术研究 Word 文档（供研究使用）。"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "research" / "Perilla_平台技术架构与应用场景研究报告.docx"


def set_doc_default_font(doc: Document, name: str = "宋体", size_pt: float = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Perilla 本地 AI 推理平台\n技术研究与应用场景分析报告")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"\n（研究稿 · {date.today().isoformat()}）\n")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph(
        "说明：本报告基于 Perilla 开源仓库公开架构文档与实现归纳，"
        "供技术调研与场景论证参考；不涉及对任何政务系统的承诺或合规认证结论。"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullet(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table_comparison(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    set_doc_default_font(doc)

    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8)
        s.right_margin = Cm(2.8)

    add_title_page(doc)

    # 摘要
    heading(doc, "摘要", 1)
    body(
        doc,
        "Perilla（紫苏）是一款面向本地与专网环境的 Web 形态 AI 推理平台。"
        "其设计哲学为「网关中心化、用户始终掌控、确定性优于魔法、本地与隐私优先」。"
        "平台将 FastAPI 推理网关作为统一控制中枢，通过 Agent Runtime、Workflow 控制面、"
        "Skill/Tool/Plugin 扩展层与 ORM/向量检索数据层，支撑对话、多步智能体、"
        "可视化工作流编排、知识库 RAG、文生图等能力。"
        "本报告重点阐述技术架构与实现要点，并深入讨论其在中国政府监管职能场景"
        "与科研领域的应用潜力，同时简要覆盖其他行业场景，并与同类开源/商业方案进行对比。",
    )
    doc.add_page_break()

    # 一、项目概述
    heading(doc, "一、项目概述与定位", 1)
    heading(doc, "1.1 产品定位", 2)
    bullet(
        doc,
        [
            "形态：Vue 3 控制台 + Python FastAPI 后端，默认本地部署（Conda 环境运行后端）。",
            "角色划分：Web UI 为「控制台」；FastAPI 为「大脑」；Agent/Plugin 为「能力模块」。",
            "推理统一：前端不直连模型或工具，所有 LLM/VLM/Embedding/ASR 经 Inference Gateway。",
            "编排统一：Workflow 为控制面产品层，与 Execution Kernel（DAG 引擎）解耦。",
        ],
    )
    heading(doc, "1.2 核心设计原则（AGENTS.md）", 2)
    bullet(
        doc,
        [
            "User-in-Control：禁止隐式代理；审批门、验收检查点等显式控制流。",
            "Gateway-Centric：模型路由、流式 SSE、错误恢复集中在网关。",
            "Determinism over Magic：不隐藏 Prompt，行为可预测、可复现。",
            "Plugin-first：能力可独立启停、替换与审计。",
            "Local-first & Privacy-first：默认无外网；数据外传需显式配置。",
        ],
    )
    heading(doc, "1.3 技术栈摘要", 2)
    table_comparison(
        doc,
        ["层次", "技术选型", "职责"],
        [
            ["前端", "Vue 3 + Vite + Tailwind + shadcn-vue", "模型/对话/Agent/工作流/知识库控制台"],
            ["API", "FastAPI", "REST、SSE 流式、OpenAPI"],
            ["Agent", "plan_based + legacy；Execution Kernel", "多步推理、Skill、重规划"],
            ["工作流", "Workflow Control Plane + GraphRuntime", "版本、执行、审批、Fork/Join、Verify Loop"],
            ["推理", "Inference Gateway + 多 Provider", "Ollama、OpenAI 兼容、llama.cpp、Torch 等"],
            ["数据", "SQLAlchemy ORM + VectorSearchProvider", "会话、知识库、工作流、租户隔离"],
        ],
    )

    # 二、技术架构
    heading(doc, "二、技术架构与实现细节", 1)
    heading(doc, "2.1 分层架构", 2)
    body(
        doc,
        "后端采用「API → Core Domain → Runtime/Agent → Skill/Tool/Plugin → Data/Infra」五层结构，"
        "并明确区分控制面与执行面：控制面负责定义、配置、路由、治理与可观测；"
        "执行面负责推理调用、工具执行、图调度与状态推进。",
    )
    bullet(
        doc,
        [
            "控制面示例：Inference Gateway、Workflow Control Plane、System Settings、治理审计。",
            "执行面示例：各类 Model Runtime、Execution Kernel、SkillExecutor、WorkflowRuntime。",
        ],
    )

    heading(doc, "2.2 推理网关（Inference Gateway, v2.8）", 2)
    body(
        doc,
        "位于 backend/core/inference/，提供统一 InferenceClient 入口，"
        "经 router 选择模型与 Provider，支持流式输出与统计指标。"
        "LLM 节点可配置 model_id 或 model_tier（low/standard/thorough），"
        "由 ModelSelector 映射到 fast/chat/reasoning 档位，便于在成本与质量间权衡。",
    )

    heading(doc, "2.3 Agent Runtime", 2)
    body(
        doc,
        "Agent 从 v1.5 legacy（Think-Action-Observation 循环）演进至 v2.x plan_based（Planner + PlanBasedExecutor + RePlan），"
        "并可选用 Execution Kernel（v2.5）将 Plan 编译为 DAG，实现可持久化、可回放、可审计的执行。"
        "v2.6 引入事件溯源；v2.8/v2.9 分别强化推理网关与运行时队列稳定。"
        "智能体通过 enabled_skills 声明能力边界，工具调用经网关 Skill 体系执行，"
        "支持 Trace 与协作会话（orchestrator_agent_id）等可观测字段。",
    )

    heading(doc, "2.4 Workflow 控制面（v3.0）", 2)
    body(
        doc,
        "Workflow 具备 Definition / Version / Execution 全生命周期："
        "编辑器（Vue Flow）产出 DAG，经 serialization 持久化为 WorkflowNode/Edge；"
        "WorkflowRuntime 与 GraphRuntimeAdapter 将编辑器语义映射为运行时节点。"
        "高价值编排能力包括：Fork/Join 并行扇出汇聚、Checkpoint 确定性验收、"
        "Verify Loop（Ralph 式迭代直至通过）、approval 人工审批门、parallel 并行限流等。"
        "ExecutionManager 提供队列、并发与配额治理；多租户下工作流与会话按 tenant_id 隔离。",
    )

    heading(doc, "2.5 数据层与知识库", 2)
    bullet(
        doc,
        [
            "关系型持久化统一走 ORM（platform.db 等），禁止业务裸 SQL。",
            "向量检索统一走 VectorSearchProvider 抽象，知识库解析/切分/嵌入/检索本地完成。",
            "会话、记忆、知识库、Agent 会话、工作流执行、治理审计等均支持租户维度。",
        ],
    )

    heading(doc, "2.6 安全、多租户与可观测", 2)
    bullet(
        doc,
        [
            "租户：TENANT_ENFORCEMENT_ENABLED 时敏感 API 前缀须 X-Tenant-Id；API Key 可绑定租户。",
            "生产护栏：SECURITY_GUARDRAILS_STRICT、DEBUG=false、RBAC/API Key scope 等。",
            "工作流：Timeline/Node Inspector、执行 call-chain、治理快照导入导出、审批 API。",
            "Agent：/agents/:id/trace、StepLog、事件流重建（调试与审计）。",
        ],
    )

    heading(doc, "2.7 前端实现要点", 2)
    body(
        doc,
        "前端为本地优先控制台：/models、/chat、/images、/agents、/workflow、/knowledge、/skills、/settings。"
        "工作流编辑器提供节点库、预检、OmX 编排模板与演示项目（Demo bundle）一键导入；"
        "运行页支持 SSE 状态流与失败报告导出。",
    )

    doc.add_page_break()

    # 三、应用场景
    heading(doc, "三、用途与应用场景", 1)
    heading(doc, "3.1 中国政府与监管职能场景（重点）", 2)
    body(
        doc,
        "政务与监管信息化普遍要求：数据不出域、过程可审计、权责清晰、人机协同而非全自动决策。"
        "Perilla 的本地部署、网关单点治理、显式审批/验收节点与多租户隔离，"
        "与上述要求具有较高契合度。以下为典型场景论证（需结合各单位等保、密评与采购规范另行评审）。",
    )

    heading(doc, "3.1.1 公文与信息发布「澄清—计划—审批—执行—验收」", 3)
    body(
        doc,
        "对应 Demo「发布简报门禁」编排：业务部门提交模糊需求（brief）→ LLM 澄清与计划模板 → "
        "approval 节点暂停待分管领导或办公厅审核 → 通过后调用 Agent/LLM 撰写正文 → "
        "Checkpoint 校验必填字段（如 text、summary）→ 输出定稿。"
        "价值：将「可发布」关口前移到人工审批与机器验收，避免模型一次性直出敏感表述。",
    )

    heading(doc, "3.1.2 监管政策多源比对与专题研判", 3)
    body(
        doc,
        "对应 Fork/Join + Verify Loop：一路 Web/法规库检索 Agent，一路内部知识库（规范性文件、历史案例）"
        "→ Join 汇聚 → 验证环迭代直至输出含 sources 的结构化报告。"
        "适用于市场监管、生态环境、应急管理等领域对「多规章交叉适用」的辅助研判"
        "（结论须人工确认，系统定位为辅助工具）。",
    )

    heading(doc, "3.1.3 行政审批与检查清单生成", 3)
    body(
        doc,
        "利用 plan_based Agent 或 Workflow 将「事项要素输入 → 条款匹配（RAG）→ 检查表生成」"
        "固化为可版本化的流程图；Checklist 节点可强制 required_keys，便于对接现有 OA 或执法系统。"
    )

    heading(doc, "3.1.4 跨委办局/多级租户隔离", 3)
    body(
        doc,
        "同一套平台实例可为不同厅局、直属单位分配 tenant_id 与 API Key 绑定，"
        "工作流定义、知识库、会话与执行历史按租户隔离，"
        "适合政务云或行业专网「平台共建、数据分域」模式。",
    )

    heading(doc, "3.1.5 合规与审计诉求", 3)
    bullet(
        doc,
        [
            "执行 Timeline 与节点级输入输出可回溯，支撑事后督查与模型应用备案材料整理。",
            "审批 API 与 approval_decisions 显式记录人机决策点。",
            "本地推理减少政务数据出境风险；外链工具（如 Web 搜索）可禁用或白名单化。",
            "注意：涉密系统须按分级保护要求选型部署，本报告不构成密评结论。",
        ],
    )

    heading(doc, "3.2 科研领域场景（重点）", 2)
    heading(doc, "3.2.1 文献与课题辅助调研", 3)
    body(
        doc,
        "课题组将 PDF/笔记入库至 Knowledge Base，经本地 Embedding 索引；"
        "并行 Agent（文献综述 + 实验设计建议）经 Join 与 Verify Loop 产出带引用的综述草稿，"
        "研究人员在 Inspector 中核对引用与实验伦理边界。",
    )

    heading(doc, "3.2.2 实验方案与项目书迭代", 3)
    body(
        doc,
        "澄清节点明确假设与变量 → 计划模板分解实验步骤 → 组会审批（approval）→ "
        "thorough 档 LLM 撰写方法学段落 → Checkpoint 确保方法、样本量、伦理说明等字段齐全。"
        "利于基金本子、伦理审查材料的结构化起草（需人工终稿）。",
    )

    heading(doc, "3.2.3 可复现的分析流水线", 3)
    body(
        doc,
        "Workflow 版本发布 + Execution Kernel 事件流，使「同一 version_id + 同一 input」"
        "可重复执行并对比节点输出；Python 节点经 python.run 工具在受控环境执行脚本，"
        "适合生物信息、社科定量等需半自动化的流程（须配合容器沙箱策略）。",
    )

    heading(doc, "3.2.4 院所私有模型与算力统一入口", 3)
    body(
        doc,
        "Inference Gateway 统一对接院内 Ollama/vLLM/商用 API，"
        "科研人员通过控制台切换 model_tier 而非散落脚本，"
        "管理方可按租户统计调用与队列并发（Runtime Stabilization）。",
    )

    heading(doc, "3.3 其他应用场景（简述）", 2)
    bullet(
        doc,
        [
            "企业知识库问答与工单分流：RAG + Agent + 可选 Workflow 审批。",
            "软件开发与代码库分析：project_intelligence、MCP/Skill 导入、多 Agent 协作。",
            "教育培训：分档 LLM 草稿+精修模板、作业反馈 Workflow。",
            "媒体与内容运营：本地文生图、多版本工作流 A/B 文案（需内容安全策略）。",
            "制造业质检辅助：VLM/Tool 节点与视觉模型插件（视部署硬件而定）。",
        ],
    )

    doc.add_page_break()

    # 四、对比分析
    heading(doc, "四、与同类项目的对比分析", 1)
    body(
        doc,
        "下表从部署形态、编排能力、Agent 模型、本地/隐私、治理等维度对比常见方案。"
        "对比基于公开资料与 Perilla 仓库实现，供研究选型参考，不代表第三方产品完整能力清单。",
    )
    table_comparison(
        doc,
        ["维度", "Perilla", "Dify / FastGPT 类", "LangFlow / Flowise", "n8n + LLM 节点", "Open WebUI + Ollama", "OmX/claw-code 类 CLI"],
        [
            ["部署", "本地/专网 Web 一体", "云 SaaS 或自托管", "自托管流式 UI", "自动化为主", "聊天为主", "CLI/IM 机器人"],
            ["推理入口", "统一 Inference Gateway", "平台内置", "节点级直连", "插件式", "Ollama 直连", "多后端脚本"],
            ["工作流", "控制面+Kernel+审批/验收", "应用/workflow 较强", "可视化强", "通用自动化强", "弱/无", "编排模板/脚本"],
            ["Agent", "plan_based+Kernel+Skill", "Agent/工具链", "Agent 组件", "有限", "简单 Agent", "多 Agent 团队概念"],
            ["人机门控", "approval+checkpoint+verify_loop", "部分有人审", "视搭建", "人工节点", "少", "偏自主"],
            ["多租户", "HTTP+ORM 租户隔离", "SaaS 成熟", "视部署", "视部署", "弱", "弱"],
            ["隐私", "Local-first 默认", "依赖部署", "依赖部署", "依赖部署", "本地友好", "本地友好"],
        ],
    )

    heading(doc, "4.1 Perilla 差异化优势", 2)
    bullet(
        doc,
        [
            "控制面与执行面清晰：Workflow 版本、发布、队列、审批与 DAG 引擎分工明确。",
            "显式编排语义：Fork/Join、Checkpoint、Verify Loop 对应监管与科研「可验证交付」诉求。",
            "不隐藏 Prompt 与参数：符合政务与科研对可解释、可审计的要求。",
            "同一平台覆盖 Chat、Agent、Workflow、KB、文生图，减少工具链碎片化。",
        ],
    )

    heading(doc, "4.2 相对短板与选型注意", 2)
    bullet(
        doc,
        [
            "生态与模板市场弱于 Dify 等成熟商业/开源社区产品。",
            "重度 SaaS 多租户计费、低代码表单等需结合 SAAS_PUBLIC_LAUNCH_GATE 等文档自行建设。",
            "CLI/IM 机器人、git worktree 多机协同等非主线能力（参考 OmX 时需知边界）。",
            "政务落地须叠加等保、密评、内容安全与模型备案等制度性要求。",
        ],
    )

    # 五、结论
    heading(doc, "五、结论与展望", 1)
    body(
        doc,
        "Perilla 将「本地优先、网关中心化、用户掌控」落实为可落地的工程架构："
        "Inference Gateway 统一推理，Agent Runtime 承担复杂推理与工具调用，"
        "Workflow 控制面提供可视化、可版本化、可审计的多步编排。"
        "在中国政府监管与科研场景中，其价值主要体现在数据域可控、人机协同门控、"
        "多源汇聚与验收闭环、多租户分域及执行可追溯；"
        "其他行业可作为私有化 AI 中台的基础组件使用。"
        "后续演进可关注：审批 UI 与工作流深度集成、子图级 Fork 实例、"
        "更细粒度 RBAC、以及与现有政务 OA/科研管理系统的标准接口对接。",
    )

    heading(doc, "参考文献与资料来源", 2)
    bullet(
        doc,
        [
            "仓库文档：docs/architecture/ARCHITECTURE.md、AGENT_ARCHITECTURE.md",
            "工作流：docs/workflow/WORKFLOW_ORCHESTRATION_OVERVIEW_ZH.md",
            "治理：AGENTS.md、docs/ops/SAAS_PUBLIC_LAUNCH_GATE_ZH.md",
            "知识库：docs/rag/KNOWLEDGE_BASE.md",
            "演示编排：demos/workflows/*.bundle.json（本地）",
        ],
    )

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    main()
